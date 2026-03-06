from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = r"C:\Users\Abhi\Desktop\s0_project\CodeQuest_Documentation.docx"

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x31, 0x6B, 0xFF)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x9B, 0x5D, 0xE5)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0xAA, 0x44)
    return p

def tbl_hdr(tbl, headers):
    row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2D2D5A')
        tcPr.append(shd)

def add_table(headers, rows_data):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl_hdr(tbl, headers)
    for row_d in rows_data:
        r = tbl.add_row()
        for i, v in enumerate(row_d):
            r.cells[i].text = str(v)
    doc.add_paragraph()
    return tbl

# ═════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('\n\n\nCodeQuest: JavaScript Warrior\n')
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x31, 0x6B, 0xFF)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('Complete Project Documentation\n')
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x9B, 0x5D, 0xE5)

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = desc_p.add_run(
    'A browser-based 2D game that teaches JavaScript through interactive coding challenges.\n'
    f'Version 1.0  |  {datetime.date.today().strftime("%B %d, %Y")}'
)
r3.font.size = Pt(11)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═════════════════════════════════════════════════════════════════════════════
h1('Table of Contents')
toc = [
    '1.  Project Overview',
    '2.  Technology Stack',
    '3.  Project Architecture & Folder Structure',
    '4.  Game Design & Mechanics',
    '5.  Full Level Curriculum (32 Levels)',
    '6.  Core Components',
    '7.  State Management (Zustand Store)',
    '8.  Sandbox — Safe Code Execution Engine',
    '9.  Sound System (Web Audio API)',
    '10. Trophy & Achievement System',
    '11. LocalStorage Save System',
    '12. How to Run & Deploy',
    '13. Pages & URL Structure',
    '14. Known Limitations & Future Roadmap',
]
for item in toc:
    p = doc.add_paragraph()
    p.add_run(item).font.size = Pt(11)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
h1('1. Project Overview')

body(
    'CodeQuest: JavaScript Warrior is an interactive 2D browser-based educational game that '
    'teaches JavaScript programming from beginner to intermediate level through real coding '
    'challenges. Instead of reading textbooks or watching passive tutorials, learners write '
    'actual JavaScript code that physically controls a warrior character in a 2D game world.'
)
doc.add_paragraph()
body(
    'When code is correct, the warrior walks, fights enemies, collects gems, casts spells, '
    'and advances to the next level. This creates immediate, visual feedback that makes '
    'learning addictive and memorable. The game covers 32 levels spanning 6 unique worlds, '
    'teaching concepts from console.log all the way to async/await, classes, and algorithms.'
)
doc.add_paragraph()

h2('Key Goals')
bullet('Teach complete JavaScript curriculum — from variables to async/await')
bullet('Make programming addictive via XP, gems, stars, daily streaks, and trophies')
bullet('Provide instant visual feedback — code runs and the world reacts in real time')
bullet('No backend, login, or internet required — runs entirely on localhost')
bullet('Track progress automatically via localStorage between browser sessions')
bullet('Accommodate different learning styles: hints, code examples, story-driven context')

doc.add_paragraph()
h2('Target Audience')
bullet('Absolute beginners learning JavaScript for the first time')
bullet('High school and college students in introductory programming courses')
bullet('Self-taught developers wanting to reinforce fundamentals')
bullet('Coding bootcamp students needing daily practice challenges')
bullet('Teachers looking for a gamified classroom tool')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ═════════════════════════════════════════════════════════════════════════════
h1('2. Technology Stack')

add_table(
    ['Category', 'Technology', 'Version', 'Purpose'],
    [
        ('Web Framework',    'Next.js',                   '14 (App Router)', 'React framework with file-based routing'),
        ('Language',         'TypeScript',                '5.x',             'Type-safe all source files'),
        ('UI Library',       'React',                     '18',              'Component-based UI rendering'),
        ('Styling',          'Vanilla CSS + Inline Styles', '—',             'Custom design system, glassmorphism, animations'),
        ('Code Editor',      'Monaco Editor',             '4.x',             'VS Code-grade in-browser JS editor'),
        ('State Management', 'Zustand',                   '4.x',             'Lightweight global state with localStorage persistence'),
        ('Code Runner',      'Function Constructor',      'Native JS',       'Sandboxed user code execution with API stubs'),
        ('Sound System',     'Web Audio API',             'Native Browser',  'Programmatic SFX — no audio files needed'),
        ('Typography',       'Google Fonts',              '—',               'Press Start 2P (pixel), Inter (UI), JetBrains Mono (code)'),
        ('Package Manager',  'npm',                       '9+',              'Dependency management'),
        ('Runtime',          'Node.js',                   '18+',             'Required for Next.js development server'),
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3. ARCHITECTURE & FOLDER STRUCTURE
# ═════════════════════════════════════════════════════════════════════════════
h1('3. Project Architecture & Folder Structure')

h2('3.1 Four-Panel Game Layout')
body(
    'The game page (/game/[level]) uses a 2x2 CSS Grid layout occupying exactly 100vh. '
    'Each panel is equal in size and has a specific dedicated function:'
)
doc.add_paragraph()
add_table(
    ['Panel', 'Position', 'Component', 'Function'],
    [
        ('Game World',       'Top-Left',     'GameScreen.tsx',       '2D game world — warrior, enemies, environment, Byte companion'),
        ('Mission Briefing', 'Top-Right',    'ExplanationPanel.tsx', 'Story + JS concept + objectives for the current level'),
        ('Code Editor',      'Bottom-Left',  'CodeEditor.tsx',       'Monaco editor + Run/Hint/Reset — where students write code'),
        ('Console Output',   'Bottom-Right', 'ConsolePanel.tsx',     'Terminal-style output from console.log and code results'),
    ]
)

h2('3.2 Folder Structure')
code_block("""game1/
+-- src/
|   +-- app/                          # Next.js App Router pages
|   |   +-- layout.tsx                # Root HTML layout, font imports
|   |   +-- globals.css               # Global CSS + all keyframe animations
|   |   +-- page.tsx                  # Home page (world map)
|   |   +-- not-found.tsx             # Custom 404 "Level Not Found" page
|   |   +-- game/
|   |   |   +-- [level]/
|   |   |       +-- page.tsx          # Game page route: /game/1-1, /game/6-3
|   |   +-- trophies/
|   |       +-- page.tsx              # Trophy room: /trophies
|   |
|   +-- components/
|   |   +-- HUD/HUD.tsx               # Top bar: XP, hearts, gems, streak, sound, trophy
|   |   +-- GameScreen/               # 2D world rendering
|   |   +-- ExplanationPanel/         # Mission briefing + story
|   |   +-- CodeEditor/CodeEditor.tsx # Monaco + run/hint/reset + code grader
|   |   +-- ConsolePanel/             # Output terminal
|   |   +-- Companion/ByteCompanion.tsx   # Floating AI orb
|   |   +-- LevelComplete/LevelCompleteModal.tsx  # Victory modal
|   |   +-- StreakCelebration/StreakCelebration.tsx
|   |
|   +-- lib/
|   |   +-- sandbox.ts                # Code execution engine
|   |   +-- sounds.ts                 # Web Audio SFX
|   |   +-- levels/
|   |       +-- types.ts              # LevelConfig interface
|   |       +-- index.ts              # ALL_LEVELS array (sorted, 32 levels)
|   |       +-- world1.ts             # Levels 1-5
|   |       +-- worlds2to4.ts         # Levels 6-20
|   |       +-- worlds5to6.ts         # Levels 22-30
|   |       +-- extra_levels.ts       # Levels 21 & 27 (gap fills)
|   |
|   +-- store/
|       +-- gameStore.ts              # Zustand global state + localStorage
|
+-- public/                           # Static assets (favicons, etc.)
+-- package.json
+-- tsconfig.json
+-- next.config.js
""")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 4. GAME DESIGN & MECHANICS
# ═════════════════════════════════════════════════════════════════════════════
h1('4. Game Design & Mechanics')

h2('4.1 Progression Flow')
numbered('Player lands on the Home page — a World Map showing all 6 worlds and 32 levels')
numbered('Levels are locked by default except Level 1-1')
numbered('Completing a level unlocks the next level in the same world')
numbered('Completing all 5-6 levels in a world unlocks World N+1')
numbered('The "Continue Quest" button routes to the first incomplete level')

doc.add_paragraph()

h2('4.2 Star Rating System')
add_table(
    ['Stars', 'Condition', 'Meaning'],
    [
        ('3 Stars', '0 hints used AND completed in under 90 seconds', 'Perfect run — efficient and unaided'),
        ('2 Stars', '1 or fewer hints AND completed in under 180 seconds', 'Good run — minor assistance'),
        ('1 Star',  'Any completion (hints or slow)',                       'Challenge cleared'),
    ]
)
body('Best star count is saved — retrying and scoring higher upgrades the stored rating.')

doc.add_paragraph()

h2('4.3 XP & Player Level')
bullet('XP is awarded on every level completion (50 XP to 500 XP based on difficulty)')
bullet('Player Level = floor(Total XP / 500) + 1')
bullet('The HUD XP bar shows progress toward next player level')
bullet('XP persists across sessions via localStorage')

doc.add_paragraph()

h2('4.4 Gems')
bullet('Gems are collectible currency awarded on level completion (1–10 per level)')
bullet('Bonus gems awarded on streak milestones: 3d=+2, 7d=+5, 14d=+10, 30d=+20')
bullet('Future use: purchasing warrior skins and power-ups')

doc.add_paragraph()

h2('4.5 Daily Streak System')
bullet('Streak tracks consecutive days the player opens the game')
bullet('Increments if: today != lastPlayed && lastPlayed == yesterday')
bullet('Resets to 1 if more than one day is missed')
bullet('Milestone celebrations at 3, 7, 14, and 30 consecutive days')
bullet('A modal appears showing streak count + bonus gems when a milestone is hit')

doc.add_paragraph()

h2('4.6 Hint System')
bullet('Every level has exactly 3 progressive hints')
bullet('Hints are revealed one at a time via the 💡 Hint button in the code editor')
bullet('Each hint reveals shows a yellow banner above the Monaco editor')
bullet('Number of hints used is tracked in the Zustand store')
bullet('"1 hint used" badge appears in the editor header after clicking Hint')
bullet('Hint usage feeds into the star calculation: 0 hints = eligible for 3 stars')

doc.add_paragraph()

h2('4.7 Code Quality Grader')
body('After a passing run, the editor shows a Code Quality score (0–100%):')
add_table(
    ['Deduction', 'Points', 'Rule'],
    [
        ('Uses var keyword',              '-15', 'Prefer let/const for safe scoping'),
        ('Uses console.error',            '-10', 'Remove debug calls from production code'),
        ('Lines longer than 100 chars',   '-10', 'Readability best practice'),
        ('No comments on code > 5 lines', '-10', 'Good code is self-documenting'),
        ('Empty unnecessary lines',       '-5',  'Code cleanliness'),
    ]
)
body('The score badge (e.g., "📊 85%") appears in the code editor header. A tip line '
     'appears below the editor with the most impactful suggestion.')

doc.add_paragraph()

h2('4.8 Byte — AI Companion')
body(
    'Byte is a floating blue orb in the top-right corner of the Game World panel. '
    'It dynamically changes color and text based on the current game phase:'
)
add_table(
    ['Phase', 'Orb Color', 'Message Examples'],
    [
        ('idle',    'Blue  (#4f8fff)', '"Read the Mission Briefing carefully..." / "Write your code and press Run!"'),
        ('running', 'Teal  (#00d4aa)', '"Processing your spell..." / "Executing..."'),
        ('success', 'Gold  (#ffd700)', '"AMAZING! Clean code!" / "You\'re a natural coder!"'),
        ('retry',   'Red   (#ff4757)', '"Hmm... check the objective again" / "Look at the console output"'),
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 5. FULL LEVEL CURRICULUM
# ═════════════════════════════════════════════════════════════════════════════
h1('5. Full Level Curriculum (32 Levels)')

worlds = [
    ('World 1: The Awakening Forest  (Theme: Teal / Forest)', [
        ('1-1',  1,  'The Sleeping Warrior',    'console.log & Output',           'Call console.log() to output text to the console'),
        ('1-2',  2,  'Name Your Warrior',       'Variables (let, const)',          'Declare let and const variables; print them'),
        ('1-3',  3,  'The Type Riddle',         'Data Types',                     'str, number, boolean, null, undefined, typeof'),
        ('1-4',  4,  'Gold Calculator',         'Arithmetic Operators',           'Use +,-,*,/,%,** in expressions; print results'),
        ('1-5',  5,  "The Elder's Scroll",      'Template Literals',              'Use backtick strings with ${} expressions'),
    ]),
    ('World 2: Village of Logic  (Theme: Orange / Village)', [
        ('2-1',  6,  'The Gatekeeper',          'Comparison & Logical Operators', '&&, ||, >=, ===, !== to evaluate entry conditions'),
        ('2-2',  7,  'Fork in the Road',        'If / Else If / Else',            'Branch code execution based on strength values'),
        ('2-3',  8,  'The Market Merchant',     'Switch Statement',               'switch/case/break/default to select purchased items'),
        ('2-4',  9,  'The Haunted Staircase',   'While Loop & Do-While',          'Climb floors using while; knock using do-while'),
        ('2-5', 10,  'The Training Gauntlet',   'For Loop, break & continue',     'for loop, skip iterations with continue, stop with break'),
    ]),
    ('World 3: The Mage Tower  (Theme: Purple / Tower)', [
        ('3-1', 11,  'The First Spell',         'Function Declaration & Calling', 'Define and call functions with parameters'),
        ('3-2', 12,  'The Return Scroll',       'Return Values',                  'Functions that return values; store in variables'),
        ('3-3', 13,  'Shadow Functions',        'Function Expressions & Arrow Fn','Write functions as expressions and arrow functions'),
        ('3-4', 14,  'The Scope Maze',          'Variable Scope',                 'Understand global, function, and block scope'),
        ('3-5', 15,  'The Echo Chamber',        'Closures & Default Parameters',  'Create closures; use default parameter values'),
    ]),
    ('World 4: The Ancient Vault  (Theme: Gold / Vault)', [
        ('4-1', 16,  'The Treasure Inventory',  'Arrays — Create, Access, Mutate','Create arrays; push/pop/shift/unshift/length'),
        ('4-2', 17,  'Sorting the Army',        'Array Methods (filter/map/reduce)','Chain .filter(), .map(), .reduce(), .find()'),
        ('4-3', 18,  "The Warrior's Profile",   'Objects — Properties & Methods', 'Create objects; dot/bracket notation; methods'),
        ('4-4', 19,  'Unpacking the Vault',     'Destructuring',                  'Destructure arrays and objects in one line'),
        ('4-5', 20,  'The Clone Army',          'Spread Operator & Rest Params',  'Spread arrays/objects; collect rest parameters into array'),
        ('4-6', 21,  'The String Sorcerer',     'String Methods',                 '.trim(), .toUpperCase(), .includes(), .split(), .slice(), .replace()'),
    ]),
    ('World 5: The Inferno Keep  (Theme: Red / Inferno)', [
        ('5-1', 22,  'Forging the Warrior Class','Classes & Constructor',          'Define a class with constructor, properties, and methods'),
        ('5-2', 23,  'The Bloodline',           'Class Inheritance & super',      'Extend a class; call super() in child constructor'),
        ('5-3', 24,  'The Cursed Code',         'try / catch / finally / throw',  'Wrap errors; throw custom errors; always-run finally block'),
        ('5-4', 25,  'The Crystal Oracle',      'Promises (.then / .catch)',       'Create and chain Promises; handle rejection gracefully'),
        ('5-5', 26,  'The Drawbridge of Time',  'Async / Await',                  'Write async functions; await sequential Promises'),
        ('5-6', 27,  'The Data Vault',          'JSON, Set & Map',                'JSON.stringify/parse; new Set(); new Map()'),
    ]),
    ('World 6: The Shadow Realm  (Theme: Violet / Shadow)', [
        ('6-1', 28,  'The Labyrinth of Logic',  'Algorithms — Nested Loops',      'Search a 2D array grid with nested for loops'),
        ('6-2', 29,  'The Mirror Mage',         'Classes + Async + Error Handling','BattleSystem class using async methods and try/catch'),
        ('6-3', 30,  'The Shadow King',         'FINAL BOSS — All Concepts',      'Combine classes, async/await, error handling, and array methods'),
    ]),
]

for world_name, levels in worlds:
    h2(world_name)
    add_table(
        ['ID', 'Level #', 'Title', 'JS Concept', 'Challenge Summary'],
        [(lid, str(ln), title, concept, challenge) for lid, ln, title, concept, challenge in levels]
    )

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 6. CORE COMPONENTS
# ═════════════════════════════════════════════════════════════════════════════
h1('6. Core Components')

h2('6.1 LevelConfig Interface (src/lib/levels/types.ts)')
body('Every level is defined by a typed object:')
code_block("""interface LevelConfig {
  id: string;            // e.g. "2-3"
  world: number;         // World 1-6
  level: number;         // Global level number 1-32
  title: string;         // "The Gatekeeper"
  worldName: string;     // "Village of Logic"
  concept: string;       // "Comparison & Logical Operators"
  story: string;         // Narrative mission text
  objective: string;     // What the code must accomplish
  starterCode: string;   // Pre-filled code in the editor
  hints: string[];       // 3 progressive hints
  validate: (logs: string[], returnValue?: unknown) => boolean;
  winMessage: string;    // Console success message
  xpReward: number;      // XP on completion
  gemReward: number;     // Gems on completion
  warriorAnimation: string; // 'walk' | 'fight' | 'magic' | 'collect' | 'celebrate'
  worldTheme: string;    // 'forest' | 'village' | 'tower' | 'vault' | 'inferno' | 'shadow'
  enemies?: { name: string; sprite: string; health: number }[];
}""")

doc.add_paragraph()

h2('6.2 HUD (src/components/HUD/HUD.tsx)')
body('The fixed top bar contains:')
bullet('CodeQuest logo and current world icon + level title')
bullet('JS Concept badge (e.g. "🧠 Template Literals")')
bullet('5 heart icons representing health (cosmetic)')
bullet('XP bar with player level number')
bullet('XP total, gem counter, day streak')
bullet('🔊/🔇 Sound toggle — toggles mute via Web Audio global flag')
bullet('🏆 Trophy Room button — navigates to /trophies')

doc.add_paragraph()

h2('6.3 CodeEditor (src/components/CodeEditor/CodeEditor.tsx)')
body('Core workflow:')
numbered('User writes/edits JavaScript in Monaco Editor')
numbered('▶ Run Code button click calls runCode(code) from sandbox.ts')
numbered('Result logs displayed in console; validate() checks for pass/fail')
numbered('Pass: completeLevel(stars, time) → saves to store + localStorage')
numbered('Pass: gradeCode(code) → shows quality % badge in header')
numbered('Fail: shows error in console; game phase = retry')
body('Additional controls:')
bullet('💡 Hint — cycles through 3 hints; tracks count in store')
bullet('🔄 Reset — restores original starterCode for the level')

doc.add_paragraph()

h2('6.4 LevelCompleteModal (src/components/LevelComplete/LevelCompleteModal.tsx)')
body('Appears after a successful level completion:')
bullet('Animated star reveal — stars appear one-by-one with sound effects')
bullet('XP and gem rewards displayed with large colorful numbers')
bullet('Three buttons: "Next Level →", "🔄 Replay", "🏠 Map"')
bullet('Accent color changes per world theme (green forest, orange village, etc.)')
bullet('Backdrop blur + slide-up entrance animation')

doc.add_paragraph()

h2('6.5 StreakCelebration (src/components/StreakCelebration/StreakCelebration.tsx)')
body('Auto-triggers when streak hits a milestone:')
add_table(
    ['Streak', 'Icon', 'Bonus Gems', 'Title'],
    [
        ('3 days',  '🔥', '+2',  'On Fire!'),
        ('7 days',  '💎', '+5',  '7-Day Streak! Crystal Armor unlocked!'),
        ('14 days', '🌟', '+10', '14-Day Streak! Legendary!'),
        ('30 days', '👑', '+20', '30-Day Streak! Crown of Code awarded!'),
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 7. STATE MANAGEMENT
# ═════════════════════════════════════════════════════════════════════════════
h1('7. State Management (Zustand Store)')

h2('7.1 Store Location')
body('File: src/store/gameStore.ts')
body('The store is created with Zustand\'s create() function. All components access state via the useGameStore() hook.')

doc.add_paragraph()

h2('7.2 State Fields')
add_table(
    ['Field', 'Type', 'Description'],
    [
        ('playerName',      'string',  'Warrior display name'),
        ('health',          'number',  'Health points (0-100, cosmetic)'),
        ('xp',              'number',  'Total experience points'),
        ('gems',            'number',  'Total gems collected'),
        ('streak',          'number',  'Consecutive day streak count'),
        ('lastPlayed',      'string',  'Date string of last play session'),
        ('currentWorld',    'number',  'Currently active world (1-6)'),
        ('currentLevel',    'number',  'Currently active level (1-32)'),
        ('levelProgress',   'Record',  'Map of level ID to { stars, completed, hintsUsed, bestTime }'),
        ('code',            'string',  'Current content of the code editor'),
        ('isRunning',       'boolean', 'True while sandbox is executing'),
        ('consoleMessages', 'array',   'Queue of console output messages'),
        ('gamePhase',       'string',  '"idle" | "running" | "success" | "retry"'),
        ('hintsUsed',       'number',  'Hints used in the current level attempt'),
        ('showHint',        'boolean', 'Whether the hint banner is visible'),
        ('currentHintIndex','number',  'Index of the currently shown hint (0,1,2)'),
    ]
)

h2('7.3 Key Actions')
add_table(
    ['Action', 'Effect'],
    [
        ('setCode(code)',            'Updates the code editor string'),
        ('addConsoleMessage(msg)',   'Appends to consoleMessages array'),
        ('clearConsole()',           'Resets consoleMessages to []'),
        ('setGamePhase(phase)',      'Switches game phase; triggers Byte companion state'),
        ('completeLevel(stars, t)',  'Saves stars/time/completed to levelProgress; calls saveToStorage()'),
        ('addXP(n)',                 'Increments xp; triggers saveToStorage()'),
        ('addGems(n)',               'Increments gems; triggers saveToStorage()'),
        ('updateStreak()',           'Checks date; increments or resets streak; saves'),
        ('toggleHint()',             'Toggles showHint; if revealing: hintsUsed++'),
        ('nextHint()',               'Increments currentHintIndex'),
        ('setLevel(world, level)',   'Resets game state for a new level'),
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 8. SANDBOX
# ═════════════════════════════════════════════════════════════════════════════
h1('8. Sandbox — Safe Code Execution Engine')

body(
    'File: src/lib/sandbox.ts\n\n'
    'The sandbox provides safe execution of user-provided JavaScript using the Function '
    'constructor approach. This isolates user code from the window object and global scope '
    'while injecting controlled game API stubs.'
)

h2('8.1 Execution Flow')
numbered('CodeEditor calls: const result = await runCode(userCode)')
numbered('A mock console object is created — captures all .log/.error/.warn calls into a string array')
numbered('Game API stub functions are defined inline (move, fight, castSpell, etc.)')
numbered('new Function(...paramNames, code) creates a sandboxed function with stubs as parameters')
numbered('fn() is called; if result is a Promise (async code), it is awaited')
numbered('Returns: RunResult { logs: string[], error: string | null, returnValue?: unknown }')
numbered('The level.validate(logs, returnValue) function checks if the challenge was solved')

doc.add_paragraph()

h2('8.2 Infinite Loop Protection')
body(
    'The execution Promise races against a 5-second timeout using Promise.race(). '
    'If the user writes while(true){} or similar, the race rejects after 5 seconds '
    'with the message: "Code timed out after 5 seconds — check for infinite loops!"'
)
code_block("""return withTimeout(execution, 5000,
  'Code timed out after 5 seconds — check for infinite loops!'
).catch(e => ({ logs: [], error: e.message }));""")

doc.add_paragraph()

h2('8.3 Game API Stubs Provided to User Code')
add_table(
    ['Stub', 'Type', 'Available for World(s)', 'Description'],
    [
        ('move(dir, steps)',         'function', '1',     'Logs warrior movement'),
        ('fight(enemy)',             'function', '1,3',   'Logs fight action'),
        ('collect(items)',           'function', '1',     'Logs item collection'),
        ('unlockDoor(code)',         'function', '1',     'Logs door unlock'),
        ('buy(item, cost)',          'function', '2',     'Logs a market purchase'),
        ('train(n)',                 'function', '2',     'Logs training action'),
        ('climbStair()',             'function', '2',     'Logs stair climbing'),
        ('castSpell(spell)',         'function', '3',     'Logs a spell cast'),
        ('takePath(path)',           'function', '2,3',   'Logs path selection'),
        ('lowerBridge(name, ms)',    'Promise',  '5',     'Resolves after delay (max 500ms)'),
        ('chargeUltimate(ms)',       'Promise',  '6',     'Resolves after charging delay'),
        ('ultimateStrike()',         'function', '6',     'Logs ultimate boss attack'),
        ('warrior',                  'Object',   '4,5',   '{ name, health, attack, walk(), cross() }'),
        ('goblin',                   'Object',   '3',     '{ name, defense, health }'),
        ('shadowKing',               'Object',   '6',     '{ defense: 80 }'),
        ('shadowArmy',               'Array',    '6',     'Array of enemy objects for final boss'),
        ('openGate(name)',           'Promise',  '6',     'Returns resolved gate name string'),
        ('pressurePlate',            'Object',   '3',     '{ onActivate: null } — user sets callback'),
        ('console',                  'Object',   'All',   'Intercepted log/error/warn/info'),
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 9. SOUND SYSTEM
# ═════════════════════════════════════════════════════════════════════════════
h1('9. Sound System (Web Audio API)')

body(
    'File: src/lib/sounds.ts\n\n'
    'All sound effects are generated programmatically using the Web Audio API. '
    'No external audio files are loaded. Waveforms (sine, square, sawtooth, triangle) '
    'are combined with gain envelopes for each effect.'
)

h2('9.1 Sound Functions')
add_table(
    ['Function', 'Trigger', 'Waveform', 'Description'],
    [
        ('sfxRun()',           'Run Code click',        'square',   'Two-tone startup blip'),
        ('sfxSuccess()',       'Validation pass',       'sine',     'Ascending 4-note melody'),
        ('sfxError()',         'Validation fail',       'sawtooth', 'Descending buzz tone'),
        ('sfxLevelComplete()', 'Victory modal open',    'sine',     'Full 7-note fanfare'),
        ('sfxHint()',          'Hint button click',     'sine',     'Soft double ping'),
        ('sfxFight()',         'Warrior fight anim',    'sawtooth', 'Low impact thud'),
        ('sfxWalk()',          'Warrior walk anim',     'square',   'Soft footstep tap'),
        ('sfxCollect()',       'Treasure collect',      'sine',     'High crystal chime'),
        ('sfxStars(n)',        'Star reveal in modal',  'sine',     'Staggered ping per star'),
        ('sfxStreak()',        'Streak milestone',      'triangle', '5-note ascending arpeggio'),
    ]
)

h2('9.2 Mute System')
body(
    'The HUD contains a 🔊/🔇 sound toggle button. Clicking it sets '
    'window.__cq_muted = true/false and saves to localStorage("codequest-muted"). '
    'Every playTone() call checks this global flag before creating any AudioContext nodes.'
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 10. TROPHY SYSTEM
# ═════════════════════════════════════════════════════════════════════════════
h1('10. Trophy & Achievement System')

body(
    'File: src/app/trophies/page.tsx\n\n'
    'The Trophy Room (/trophies) displays all 16 achievements evaluated dynamically '
    'against the Zustand game state. No IDs are stored in localStorage — conditions '
    'are recalculated live each visit.'
)

h2('10.1 Achievement List')
add_table(
    ['Icon', 'Title', 'Rarity', 'Unlock Condition'],
    [
        ('⚔️', 'First Blood',      'COMMON',    'Complete any 1 level'),
        ('🧠', 'Big Brain',        'COMMON',    'Complete a level with 0 hints used'),
        ('⭐', 'Perfectionist',    'RARE',      'Earn 3 stars on any level'),
        ('🌲', 'Forest Cleared',   'RARE',      'Complete all levels in World 1'),
        ('🏰', 'Logic Master',     'RARE',      'Complete all levels in World 2'),
        ('🔮', 'Spell Weaver',     'EPIC',      'Complete all levels in World 3'),
        ('🏺', 'Vault Raider',     'EPIC',      'Complete all levels in World 4'),
        ('🌋', 'Inferno Walker',   'EPIC',      'Complete all levels in World 5'),
        ('🏆', 'JS Champion',      'LEGENDARY', 'Complete all 32 levels'),
        ('🔥', 'On Fire',          'COMMON',    'Maintain 3-day streak'),
        ('💎', 'Diamond Habit',    'EPIC',      'Maintain 7-day streak'),
        ('✨', 'XP Farmer',        'COMMON',    'Earn 500 total XP'),
        ('🌟', 'XP Legend',        'RARE',      'Earn 2000 total XP'),
        ('💍', 'Gem Collector',    'COMMON',    'Collect 10 gems'),
        ('⚡', 'Speed Coder',      'RARE',      'Complete a level in under 30 seconds'),
        ('🌠', 'Star Gazer',       'LEGENDARY', 'Earn 3 stars on 5 different levels'),
    ]
)

h2('10.2 Rarity Color System')
add_table(
    ['Rarity', 'Background Color', 'Border Color', 'Text Color'],
    [
        ('COMMON',    '#1a2a1a (dark green)',  '#3a5a3a (green)',         '#7aba7a (light green)'),
        ('RARE',      '#0d1a2e (dark blue)',   '#2a5090 (blue)',          '#6aaaff (light blue)'),
        ('EPIC',      '#1a0d2e (dark purple)', '#7a3aad (purple)',        '#c07aff (light purple)'),
        ('LEGENDARY', '#2e1a00 (dark gold)',   '#b8860b (dark goldenrod)','#ffd700 (gold)'),
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 11. LOCALSTORAGE SAVE
# ═════════════════════════════════════════════════════════════════════════════
h1('11. LocalStorage Save System')

h2('11.1 Save Keys')
add_table(
    ['Key', 'Type', 'Content'],
    [
        ('"codequest-save"',  'JSON string', 'All player progress: XP, gems, streak, level completions, stars'),
        ('"codequest-muted"', 'string',      '"true" or "false" — audio mute preference'),
    ]
)

h2('11.2 Save Data Structure')
code_block("""{
  "playerName": "Warrior",
  "xp": 525,
  "gems": 12,
  "streak": 4,
  "lastPlayed": "Wed Mar 04 2026",
  "currentWorld": 2,
  "currentLevel": 8,
  "levelProgress": {
    "1-1": { "stars": 3, "completed": true, "hintsUsed": 0, "bestTime": 42 },
    "1-2": { "stars": 2, "completed": true, "hintsUsed": 1, "bestTime": 98 },
    "1-3": { "stars": 3, "completed": true, "hintsUsed": 0, "bestTime": 61 }
  }
}""")

h2('11.3 Reset Progress')
body('To clear all saved data, run in the browser DevTools Console (F12):')
code_block('localStorage.removeItem("codequest-save"); location.reload();')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 12. HOW TO RUN & DEPLOY
# ═════════════════════════════════════════════════════════════════════════════
h1('12. How to Run & Deploy')

h2('12.1 Prerequisites')
add_table(
    ['Requirement', 'Minimum Version', 'Notes'],
    [
        ('Node.js',  '18.0',     'Required for Next.js 14'),
        ('npm',      '9.0',      'Comes with Node.js'),
        ('Browser',  'Chrome 90 / Firefox 88 / Edge 90', 'Web Audio API required for sounds'),
        ('RAM',      '4GB',      'Monaco Editor can be memory-hungry'),
        ('Internet', 'Optional', 'Only needed to load Google Fonts on first visit'),
    ]
)

h2('12.2 Development Setup')
numbered('Clone or download the project to your machine')
numbered('Open a terminal in the project directory:')
code_block('cd C:\\Users\\Abhi\\Desktop\\s0_project\\game1')
numbered('Install dependencies:')
code_block('npm install')
numbered('Start the dev server:')
code_block('npm run dev')
numbered('Open browser at: http://localhost:3000')

doc.add_paragraph()

h2('12.3 Production Build')
code_block('npm run build   # Compile optimized production bundle\nnpm start       # Start production server on port 3000')

doc.add_paragraph()

h2('12.4 Environment Notes')
bullet('No .env file is required — the app has no external API dependencies')
bullet('No database — localStorage is the complete persistence layer')
bullet('The dev server supports Hot Module Replacement (HMR) for instant code changes')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 13. PAGES & URL STRUCTURE
# ═════════════════════════════════════════════════════════════════════════════
h1('13. Pages & URL Structure')

add_table(
    ['URL', 'File', 'Description'],
    [
        ('/',          'src/app/page.tsx',                  'Home — World Map, player stats, level selection, Continue Quest button'),
        ('/game/1-1',  'src/app/game/[level]/page.tsx',     'Level 1-1: The Sleeping Warrior (console.log)'),
        ('/game/2-3',  'src/app/game/[level]/page.tsx',     'Level 2-3: The Market Merchant (switch statement)'),
        ('/game/4-6',  'src/app/game/[level]/page.tsx',     'Level 4-6: The String Sorcerer (string methods)'),
        ('/game/5-6',  'src/app/game/[level]/page.tsx',     'Level 5-6: The Data Vault (JSON, Set, Map)'),
        ('/game/6-3',  'src/app/game/[level]/page.tsx',     'Level 6-3: Final Boss — The Shadow King'),
        ('/trophies',  'src/app/trophies/page.tsx',         'Trophy Room — all 16 achievements, progress bar, filter tabs'),
        ('/*',         'src/app/not-found.tsx',             'Custom 404 — "Level Not Found" with Back to Map button'),
    ]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 14. LIMITATIONS & FUTURE ROADMAP
# ═════════════════════════════════════════════════════════════════════════════
h1('14. Known Limitations & Future Roadmap')

h2('14.1 Current Limitations')
add_table(
    ['Issue', 'Impact', 'Proposed Fix'],
    [
        ('Synchronous infinite loops block main thread for 5s',   'Medium', 'Move execution to a Web Worker thread'),
        ('No multiplayer or leaderboard',                          'Low',    'Add Supabase or Firebase backend'),
        ('4-panel layout cramped on screens < 900px',              'Medium', 'Add responsive breakpoints — vertical stack on mobile'),
        ('Monaco shows a minor ESLint badge in some environments', 'Low',    'Suppress specific rules via Monaco editor config'),
        ('Web Audio needs a user interaction first on Chrome',      'Low',    'AudioContext.resume() on first click — already handled'),
        ('World progress indicator in HUD hardcoded to 30',        'Low',    'Use ALL_LEVELS.length dynamically'),
    ]
)

h2('14.2 Future Roadmap')
add_table(
    ['Feature', 'Priority', 'Description'],
    [
        ('Web Worker Sandbox',            'High',   'Fully isolate code execution off main thread; true infinite loop immunity'),
        ('Mobile Responsive Layout',      'High',   'Stack panels vertically on screens < 768px'),
        ('World 7: DOM & Browser APIs',   'High',   '5 levels teaching document.querySelector, events, fetch, localStorage'),
        ('World 8: Node.js Basics',       'Medium', 'fs, http, modules, npm basics in a simulated terminal'),
        ('Warrior Skins Shop',            'Medium', 'Spend gems on cosmetic skins unlocked via achievements'),
        ('Phaser.js Game Engine',         'Medium', 'Replace CSS 2D world with real tile-based Phaser scene'),
        ('Teacher Dashboard',             'Medium', 'Admin view to see student completion stats and star ratings'),
        ('Cloud Save (Optional Login)',   'Medium', 'Firebase Auth + Firestore for cross-device progress'),
        ('World Cutscene Transitions',    'Low',    '3-second cinematic when entering a new world for the first time'),
        ('Ghost Replay Mode',             'Low',    'Watch the highest-starred code solution after completing a level'),
        ('AI Hint Generator',             'Low',    'Use Gemini API to generate contextual hints based on user code'),
        ('Leaderboard',                   'Low',    'Optional name entry; rank by total XP or speedrun times'),
    ]
)

doc.add_paragraph()

# ─ Final separator ─
doc.add_paragraph('─' * 80)
fin = doc.add_paragraph()
fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_fin = fin.add_run(
    f'CodeQuest: JavaScript Warrior  |  Project Documentation  |  '
    f'Generated {datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p")}'
)
r_fin.font.size = Pt(9)
r_fin.font.color.rgb = RGBColor(0x88, 0x88, 0xAA)

# ─ Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"SUCCESS: Document saved to {OUTPUT}")
